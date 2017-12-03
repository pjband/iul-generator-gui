# -*- coding: utf-8 -*-

import os
import time
import hashlib
import configparser
import docx
from docx.shared import Pt
import copy
import shutil
import re
import traceback
import logging
from tkinter import *
from tkinter import ttk
from tkinter.filedialog import askdirectory
from tkinter import messagebox
from tkinter import filedialog


logging.basicConfig(filename='errorlog.log',
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                    level=logging.INFO)

lang = {
    'byte': 'Б',
    'file': 'Файл',
    'changed': 'Изменен',
    'size': 'Размер',
    'checksum_sha1': 'Контр. сумма SHA-1',
    'list_of_documents': 'Ведомость электронных документов',
    'identification_sheet': 'ИУЛ',
    'identification_sheet_title': 'ВЕДОМОСТЬ ЭЛЕКТРОННЫХ ДОКУМЕНТОВ',

    'directory': 'Директория',
    'wrong_directory': 'Несуществующая директория',
    'creator': 'Разработчик ИУЛ',
    'normcontroller': 'Нормоконтролер',
    'project_name': 'Номер проета',
    'result': 'Результат',
    'iul_msg': 'Не забудьте проверить правильность наименований ИУЛ',
    'iul_creation_fail': 'Что-то пошло не так. Возможно, указана несуществующая директория, ' \
              'или у вас отсутствуют права на запись в эту директорию. Детали ошибки в log-файле',

}

os.environ['TZ'] = 'Europe/Moscow'
cfg = dict()

config = configparser.ConfigParser()
config.read('config.ini', encoding='utf8')
try:
    cfg['exclude_types'] = config['exclude']['types']
    cfg['exclude_types'] = tuple(map(str.strip, cfg['exclude_types'].split(',')))
    cfg['default_iul_name'] = config['default']['iul_name']
    cfg['file_list_txt'] = config['default']['file_list_txt'].strip()
except:
    logging.exception("message")
    cfg['exclude_types'] = ()
    cfg['default_iul_name'] = 'Новый-УЛ.docx'
    cfg['file_list_txt'] = lang['list_of_documents'] + '.txt'

cfg['clearFile'] = 'clear.docx'
cfg['iul_template'] = 'iul_template.docx'
# cfg['newIUL'] = lang['identification_sheet']
cfg['fileTitle'] = lang['identification_sheet_title'] + '\r\n\r\n' \
                   + lang['file'] + ': ' + cfg['file_list_txt'] + '\r\n   ' \
                   + lang['changed'] + ': ' + str(time.strftime("%d.%m.%y %X", time.localtime())) + '\r\n\r\n'

cfg['signatories'] = ['creator', 'reviewer', 't_controller', 'normcontroller', 'accepter']

def disable_widget(widget):
    widget.config(state="disabled")


def templates_checked():
    if os.path.exists(cfg['iul_template']) and os.path.exists(cfg['clearFile']):
        return True
    return False

def enable_widget(widget):
    widget.config(state="normal")


def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = copy.deepcopy(tbl)
    p.addnext(new_tbl)


def insert_date():
    return time.strftime("%d.%m.%y", time.localtime())


def add_table(template, input_dict, clearFile):
    iul = docx.Document(clearFile)
    style = iul.styles['Normal']
    font = style.font
    paragraph_format = iul.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.left_indent = Pt(0)
    paragraph_format.first_line_indent = Pt(0)
    font.size = Pt(10)
    font.name = 'Times New Roman'

    tmpl = docx.Document(template)
    tpl = tmpl.tables[0]
    r = 0
    for row in tpl.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for key in input_dict.keys():

                    if paragraph.text in cfg['signatories'] and len(input_dict[paragraph.text]) > 0:
                        tpl.cell(r, 7).text = insert_date()

                    paragraph.text = paragraph.text.replace(str(key), str(input_dict[key]))
                    paragraph.style = iul.styles['Normal']
        r += 1
    # add table in the end
    copy_table_after(tpl, iul.paragraphs[-1])
    iul.add_paragraph('')
    # page break
    if len(iul.tables) % 2 == 0:
        iul.add_page_break()
    iul.save(clearFile)


def input_dir():
    r = {}
    r['dir'] = str(input(lang['directory'] + ': '))
    if os.path.exists(r['dir']) is False:
        return False
    else:
        r['creator'] = str(input(lang['creator'] + ': '))
        r['normcontroller'] = str(input(lang['normcontroller'] + ': '))
        return r


def generate(root_dir, creator, controller, file_abs_path, reviewer, t_control, approve):
    inp={}
    inp['creator'] = creator
    inp['normcontroller'] = controller
    repl = {}
    i = 1
    new_iul_file_docx = file_abs_path  # os.path.join(save_path, cfg['newIUL'])
    iul_filename = os.path.basename(new_iul_file_docx)

    try:
        shutil.copy(cfg['clearFile'], new_iul_file_docx)
    except:
        logging.exception("message")
        return False

    # list of documents:
    new_file_txt = os.path.join(os.path.dirname(new_iul_file_docx), cfg['file_list_txt'])
    with open(new_file_txt, 'w', encoding='utf8') as f:
        f.write(cfg['fileTitle'])
        for top, dirs, nondirs in os.walk(root_dir):
            # if os.path.basename(top) != cfg['exclude']:
            for name in nondirs:
                path = str(os.path.join(top, name))
                if path.lower().endswith(cfg['exclude_types']) is False and \
                        name != iul_filename and \
                        name != cfg['file_list_txt']:
                    filename = os.path.basename(path)
                    modTime = time.strftime("%d.%m.%y %X", time.localtime(os.stat(path).st_mtime))
                    size = os.stat(path).st_size
                    try:
                        sha1 = hashlib.sha1(open(path, 'rb').read()).hexdigest()
                        f.write(
                            lang['file'] + ': ' + str(filename) + '\r\n' +
                            '   ' + lang['changed'] + ': ' + str(modTime) + '\r\n' +
                            '   ' + lang['size'] + ': ' + str(size) + ' ' + lang['byte'] + '\r\n' +
                            '   ' + lang['checksum_sha1'] + ': ' + str(sha1) + '\r\n\r\n'
                        )
                        # changing identification sheet
                        repl['Npp'] = str(i)
                        i = i + 1
                        repl['file'] = filename
                        repl['sha-1_hash'] = sha1
                        repl['creator'] = inp['creator']
                        repl['normcontroller'] = inp['normcontroller']
                        repl['reviewer'] = reviewer
                        repl['t_controller'] = t_control
                        repl['accepter'] = approve
                        repl['iul_name'] = iul_name(filename)
                        # repl['date'] = time.strftime("%d.%m.%y", time.localtime())
                        repl['num'] = str(i // 2)
                        repl['count'] = ''
                        add_table(cfg['iul_template'], repl, new_iul_file_docx)
                        set_status(i)
                    except:
                        logging.exception("message")
                        return False
    return True


def clear_form():
    tkDIR.delete(0, 'end')
    tkCREATOR.delete(0, 'end')
    tkNORMCONTROL.delete(0, 'end')


def iul_name(filename):
    iul_name = re.search(r'\d{1,3}(.\d{1,3})?-\d{4}-\w{1,3}', filename)  # iul_name = re.search(r'\d{2}-\d{4}-\w{1,3}', filename)
    if iul_name is None:
        return filename
    else:
        return iul_name.group(0)


def open_dir():
    name = askdirectory(title="Директория")
    if name != '' and not isinstance(name, tuple):
        if os.access(name, os.W_OK):
            tkDIR.delete(0, 'end')
            tkDIR.insert(END, name)
        else:
            messagebox.showinfo('Права доступа', 'У Вас отсутствуют права на запись в указанную директорию')


def set_status(val):

    def progress_bar(i):
        i = int(i)
        root.update_idletasks()
        tkPrgs['value'] = i

    val = int(val)
    if 0 < val < 49:
        percentage = 2 * val
    elif val >= 49:
        percentage = 98
    else:
        percentage = 100

    progress_bar(percentage)


def iul_button0():
    save_path = askdirectory(initialdir=tkDIR.get(), title="Сохранить в")
    #save_path = filedialog.asksaveasfilename(initialdir=tkDIR.get(), initialfile='Новый-УЛ.docx', title="Select file", filetypes=(("docx", "*.docx"), ("Все файлы", "*.*")))
    if save_path != '' and not isinstance(save_path, tuple):
        disable_widget(tkGenBut)
        tkPrgs.grid(row=3, column=1, sticky='NE', padx=5, pady=17)
        r = generate(tkDIR.get(), tkCREATOR.get(), tkNORMCONTROL.get(), save_path)
        if r:
            mes = cfg['newIUL'] + ' успешно создан в ' + save_path
            set_status(0)
        else:
            mes = 'Что-то пошло не так. Возможно, указана несуществующая директория, ' \
                  'или у вас отсутствую права на запись в эту директорию. Детали ошибки в log-файле'
        enable_widget(tkGenBut)
        tkPrgs.grid_remove()
        messagebox.showinfo('Результат', mes)


def iul_button():
    new_file_name = filedialog.asksaveasfilename(
        initialdir=tkDIR.get(),
        initialfile=cfg['default_iul_name'],
        title="Название файла",
        filetypes=(("docx", "*.docx"), ("Все файлы", "*.*")))

    if templates_checked():

        if new_file_name != '' and not isinstance(new_file_name, tuple):

            name, ext = os.path.splitext(new_file_name)
            if ext != '.docx':
                new_file_name += '.docx'

            disable_widget(tkGenBut)
            tkPrgs.grid(row=7, column=1, sticky='NE', padx=5, pady=17)

            r = generate(tkDIR.get(),
                         tkCREATOR.get(),
                         tkNORMCONTROL.get(),
                         new_file_name,
                         tkREVIEWER.get(),
                         tkTCONTROL.get(),
                         tkAPPROVE.get())
            if r:
                mes = lang['identification_sheet'] + ' успешно создан: ' + new_file_name
                set_status(0)
            else:
                mes = lang['iul_creation_fail']

            enable_widget(tkGenBut)
            tkPrgs.grid_remove()
            messagebox.showinfo('Результат', mes)

    else:
        messagebox.showinfo('Ошибка', 'В директории программы отсутствует один из файлов-шаблонов: '
                            + cfg['iul_template'] + ' или '+cfg['clearFile']
                            + '. Восстановите эти файлы из исходного архива программы.')


if __name__ == "__main__":

    root = Tk()
    Title = root.title("ИУЛ Генератор v1.0")

    label = ttk.Label(root, text="Директория с файлами:").grid(row=0, column=0, sticky='NE', pady=4)
    tkDIRbutton = ttk.Button(root, text="Обзор..", command=open_dir).grid(row=0, column=2, sticky='NW', padx=5, pady=3)
    tkDIR = ttk.Entry(root, width=20)
    tkDIR.grid(row=0, column=1, sticky='NW', pady=5)

    label2 = ttk.Label(root, text="Разработал:").grid(row=1, column=0, sticky='NE', pady=2)
    tkCREATOR = ttk.Entry(root, width=20)
    tkCREATOR.grid(row=1, column=1, sticky='NW', pady=2)

    label22 = ttk.Label(root, text="Проверил:").grid(row=2, column=0, sticky='NE', pady=2)
    tkREVIEWER = ttk.Entry(root, width=20)
    tkREVIEWER.grid(row=2, column=1, sticky='NW', pady=2)

    label23 = ttk.Label(root, text="Т.контроль:").grid(row=3, column=0, sticky='NE', pady=2)
    tkTCONTROL = ttk.Entry(root, width=20)
    tkTCONTROL.grid(row=3, column=1, sticky='NW', pady=2)

    label3 = ttk.Label(root, text="Н.контроль:").grid(row=4, column=0, sticky='NE', pady=2)
    tkNORMCONTROL = ttk.Entry(root, width=20)
    tkNORMCONTROL.grid(row=4, column=1, sticky='NW', pady=2)

    label5 = ttk.Label(root, text="Утвердил:").grid(row=5, column=0, sticky='NE', pady=2)
    tkAPPROVE = ttk.Entry(root, width=20)
    tkAPPROVE.grid(row=5, column=1, sticky='NW', pady=2)

    tkPrgs = ttk.Progressbar(root, orient=HORIZONTAL, length=100, mode='determinate')
    tkPrgs.grid(row=7, column=1, sticky='NE', padx=5, pady=17)
    tkPrgs.grid_remove()

    labelTitle = ttk.Label(root, text="PjBand.ru").grid(row=8, column=0, sticky='NW', padx=5)

    tkGenBut = ttk.Button(root, text="Создать ИУЛ", command=iul_button)
    tkGenBut.grid(row=7, column=2, sticky='NW', padx=5, pady=10)

    root.mainloop()
